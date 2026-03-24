from flask import render_template, request, Blueprint, jsonify, session, current_app
import infrastructure.model.MCasas as MCasas
import domain.VCasas as VCasas
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from domain.VPermisos import requiere_permiso
import os
from werkzeug.utils import secure_filename
from fpdf import FPDF
from datetime import datetime
from flask import send_file
import io

bp = Blueprint('RCasas', __name__)

UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                             '..', '..', 'presentation', 'static', 'uploads', 'logos')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp', 'svg'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ── Vistas ─────────────────────────────────────────────

@bp.route('/casas')
def casas():
    return render_template('Casas/Casas.html', active_page='casas')

@bp.route('/casa/<casa_id>')
def detalle_casa(casa_id):
    return render_template('Casas/Casas.html', casa_id=casa_id, active_page='casas')

# ── API JSON ────────────────────────────────────────────

@bp.route('/get_casas', methods=['GET'])
def get_casas():
    try:
        query = request.args.get('q', '').strip()
        tipo = request.args.get('tipo', 'todos').strip()
        casas_list = MCasas.getAllCasas(query if query else None, tipo)
        
        result = [{
            '_id': str(c['_id']),
            'nombre': c.get('nombre', ''),
            'obras': c.get('obras', []),
            'historia': c.get('historia', ''),
            'tipo': c.get('tipo', 'masculino')
        } for c in casas_list]
        
        return jsonify({"success": True, "casas": result})
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc)}), 500

@bp.route('/get_casa/<casa_id>', methods=['GET'])
def get_casa(casa_id):
    try:
        casa_data = MCasas.getCasaById(casa_id)
        validator = VCasas.getCasaValidator(casa_data)
        casa = validator.validation()
        return jsonify({
            "success": True,
            "casa": {
                '_id': str(casa['_id']),
                'nombre': casa.get('nombre', ''),
                'obras': casa.get('obras', []),
                'historia': casa.get('historia', ''),
                'tipo': casa.get('tipo', 'masculino')
            }
        })
    except LookupError as exc:
        return jsonify({"success": False, "message": str(exc)}), 404
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc)}), 500

@bp.route('/upload_logo', methods=['POST'])
@requiere_permiso('casas', 'crear')
def upload_logo():
    if 'logo' not in request.files:
        return jsonify({"success": False, "message": "No se envió archivo"}), 400
    file = request.files['logo']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({"success": False, "message": "Archivo no válido"}), 400
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    filename = secure_filename(file.filename)
    file.save(os.path.join(UPLOAD_FOLDER, filename))
    return jsonify({"success": True, "filename": filename})

@bp.route('/create_casa', methods=['POST'])
@requiere_permiso('casas', 'crear')
def create_casa():
    try:
        data = request.get_json(silent=True)
        validator = VCasas.createCasaValidator(isJson=request.is_json, payLoad=data)
        casa_data = validator.validation()
        result = MCasas.createCasa(casa_data)
        return jsonify({
            "success": True,
            "message": "Casa Salesiana creada exitosamente",
            "casa_id": str(result.inserted_id)
        }), 201
    except ValueError as exc:
        return jsonify({"success": False, "message": str(exc)}), 400
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc)}), 500

@bp.route('/update_casa/<casa_id>', methods=['PUT'])
@requiere_permiso('casas', 'editar')
def update_casa(casa_id):
    try:
        data = request.get_json(silent=True)
        validator = VCasas.updateCasaValidator(isJson=request.is_json, payLoad=data)
        casa_data = validator.validation()
        result = MCasas.updateCasa(casa_id, casa_data)
        if result.modified_count > 0:
            return jsonify({"success": True, "message": "Casa actualizada exitosamente"})
        return jsonify({"success": False, "message": "No se realizaron cambios"}), 200
    except ValueError as exc:
        return jsonify({"success": False, "message": str(exc)}), 400
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc)}), 500

@bp.route('/delete_casa/<casa_id>', methods=['DELETE'])
@requiere_permiso('casas', 'eliminar')
def delete_casa(casa_id):
    try:
        result = MCasas.deleteCasa(casa_id)
        if result.deleted_count > 0:
            return jsonify({"success": True, "message": "Casa eliminada exitosamente"})
        return jsonify({"success": False, "message": "Casa no encontrada"}), 404
    except Exception as exc:
        return jsonify({"success": False, "message": str(exc)}), 500

class PDF(FPDF):
    def header(self):
        import os
        # Logo Image
        logo_path = os.path.join(current_app.root_path, 'presentation', 'static', 'img', 'logo_sdb.png')
        if os.path.exists(logo_path):
            self.image(logo_path, 10, 8, 20)
            
        # Institutional Text shifted to the right
        self.set_y(12)
        self.set_x(33)
        self.set_font("helvetica", "B", 12)
        self.set_text_color(44, 62, 80) # SDB Blue
        self.cell(0, 6, "CONDOR, La Vega R.D.", ln=True, align='L')
        
        self.set_x(33)
        self.set_font("helvetica", "", 10)
        self.cell(0, 5, "Diocesis de La Vega", ln=True, align='L')
        
        # Right-aligned Title
        self.set_y(12)
        self.set_font("helvetica", "B", 14)
        self.set_text_color(220, 30, 70) # SDB Red
        self.cell(0, 15, "DIRECTORIO DE OBRAS", ln=True, align='R')
        
        # Line
        self.set_draw_color(220, 30, 70)
        self.set_line_width(0.8)
        self.line(10, 28, 200, 28)
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Página {self.page_no()}/{{nb}}", align='C')
        self.cell(0, 10, "Agenda SDB - Sistema de Gestión", align='R')

@bp.route('/reporte_casas')
def reporte_casas():
    try:
        casas_list_raw = MCasas.getAllCasas()
        
        # 1. Agrupar casas por tipo (Solo Masculino y Femenino como pidió el usuario)
        labels = {
            'masculino': 'CASAS MASCULINAS (SDB)',
            'femenino': 'CASAS FEMENINAS (FMA / HH.SS.)'
        }
        
        grupos = {
            'masculino': [],
            'femenino': []
        }
        
        for c in casas_list_raw:
            t = str(c.get('tipo', 'masculino')).lower()
            if t == 'femenino':
                grupos['femenino'].append(c)
            else:
                grupos['masculino'].append(c)
        
        # Create PDF instance
        pdf = PDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        
        # Report Date
        pdf.set_font("helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        now = datetime.now().strftime("%d/%m/%Y %H:%M")
        pdf.cell(0, 8, f"Fecha de emisión: {now}", ln=True, align='R')
        pdf.ln(2)
        
        # Table Header Configuration
        cols = ["Nombre de la Obra", "Ciudad", "Teléfono", "Contacto"]
        w = [60, 35, 40, 55] # Widths for columns
        
        # 2. Iterar por grupos
        first_group = True
        for tipo, casas in grupos.items():
            if not casas:
                continue
                
            # Siempre iniciar un nuevo grupo en una página nueva para dividir claramente
            if not first_group:
                pdf.add_page()
            first_group = False
            
            # Título de Sección (Grupo)
            pdf.set_font("helvetica", "B", 16)
            pdf.set_text_color(220, 30, 70) # SDB Red
            pdf.cell(0, 12, labels.get(tipo, tipo.upper()), ln=True, align='C')
            pdf.set_draw_color(220, 30, 70)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(8)
            
            for c in casas:
                nombre_casa = c.get('nombre', '').strip()
                obras = c.get('obras', [])
                
                # --- LÓGICA DE PAGINACIÓN SMART ---
                # Calcular altura necesaria: 
                # Cabecera Casa(9) + Subcabecera(8) + (Num Obras * 8) + Margen(5)
                # Si no hay obras, cuenta como 1 fila de "Sin obras"
                num_filas = len(obras) if obras else 1
                altura_bloque = 9 + 8 + (num_filas * 8) + 5
                
                # Si el bloque supera el espacio disponible (trigger en ~270mm para A4)
                if pdf.get_y() + altura_bloque > 275:
                    pdf.add_page()
                
                # 3. Casa Header Row (Full Width)
                pdf.set_font("helvetica", "B", 10)
                pdf.set_fill_color(44, 62, 80) # SDB Blue
                pdf.set_text_color(255, 255, 255)
                pdf.cell(sum(w), 9, f" CASA: {nombre_casa.upper()}", border=1, ln=True, fill=True, align='L')
                
                # 4. Sub-header for Columns
                pdf.set_font("helvetica", "B", 9)
                pdf.set_fill_color(220, 30, 70) # SDB Red
                pdf.set_text_color(255, 255, 255)
                for i in range(len(cols)):
                    pdf.cell(w[i], 8, cols[i], border=1, fill=True, align='C')
                pdf.ln()
                
                # 5. Obras Rows
                pdf.set_font("helvetica", "", 9)
                pdf.set_text_color(0, 0, 0)
                
                if not obras:
                    pdf.set_fill_color(255, 255, 255)
                    pdf.cell(w[0], 8, " (Sin obras registradas)", border=1, fill=True)
                    pdf.cell(w[1], 8, "—", border=1, fill=True, align='C')
                    pdf.cell(w[2], 8, "—", border=1, fill=True, align='C')
                    pdf.cell(w[3], 8, "—", border=1, fill=True, align='C')
                    pdf.ln()
                else:
                    fill = False
                    for o in obras:
                        # Zebra striping
                        pdf.set_fill_color(248, 249, 250) if fill else pdf.set_fill_color(255, 255, 255)
                        
                        nombre_obra = (o.get('nombre_obra') or '—')[:40]
                        ap = o.get('apartado_postal')
                        ciudad_text = o.get('ciudad') or '—'
                        if ap:
                            ciudad_text = f"{ciudad_text} (AP: {ap})"
                        ciudad = str(ciudad_text)[:20]
                        
                        telfs = o.get('telefono', [])
                        tel = ", ".join([str(t) for t in telfs])[:20] if isinstance(telfs, list) else str(telfs)[:20]
                        
                        correos = o.get('correo', [])
                        cont_val = (o.get('contacto') or '—')
                        if correos:
                            email_str = ", ".join(correos) if isinstance(correos, list) else str(correos)
                            cont = f"{cont_val} ({email_str})"[:30]
                        else:
                            cont = cont_val[:30]
                        
                        pdf.cell(w[0], 8, f" {nombre_obra}", border=1, fill=True)
                        pdf.cell(w[1], 8, ciudad, border=1, fill=True, align='C')
                        pdf.cell(w[2], 8, tel, border=1, fill=True, align='C')
                        pdf.cell(w[3], 8, cont, border=1, fill=True)
                        pdf.ln()
                        fill = not fill
                
                pdf.ln(4) # Space between Casas
            
        # Output to buffer
        # Output PDF as a stream
        output = io.BytesIO()
        pdf_output = pdf.output()
        output.write(pdf_output)
        output.seek(0)
        
        filename = f"Reporte_Casas_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as exc:
        print(f"Error generando reporte: {exc}")
        return "Error al generar el reporte", 500

@bp.route('/reporte_casa/<casa_id>')
def reporte_casa(casa_id):
    try:
        casa = MCasas.getCasaById(casa_id)
        if not casa:
            return "Casa no encontrada", 404
            
        # Create PDF instance
        pdf = PDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        
        # Report Date
        pdf.set_font("helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        now = datetime.now().strftime("%d/%m/%Y %H:%M")
        pdf.cell(0, 8, f"Ficha generada el: {now}", ln=True, align='R')
        pdf.ln(5)
        
        # Title
        pdf.set_font("helvetica", "B", 16)
        pdf.set_text_color(44, 62, 80)
        pdf.cell(0, 10, casa.get('nombre', 'Detalle de la Casa'), ln=True, align='C')
        pdf.ln(10)
        
        # History section
        historia = casa.get('historia', '').strip()
        if historia:
            pdf.set_font("helvetica", "B", 12)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 10, "  Reseña Histórica", ln=True, fill=True)
            pdf.ln(4)
            pdf.set_font("helvetica", "", 10)
            pdf.multi_cell(0, 6, historia)
            pdf.ln(5)

        # Obras
        obras = casa.get('obras', [])
        if obras:
            pdf.set_font("helvetica", "B", 14)
            pdf.set_text_color(220, 30, 70)
            pdf.cell(0, 10, "Obras Asociadas", ln=True)
            pdf.set_text_color(0, 0, 0)
            
            for o in obras:
                detalles_obra = [
                    ("Ciudad:", o.get('ciudad', '—')),
                    ("Teléfono:", o.get('telefono', '—')),
                    ("Dirección:", o.get('direccion', '—')),
                    ("Sitio Web:", o.get('web', '—')),
                    ("Correo(s):", ", ".join(o.get('correo', [])) if isinstance(o.get('correo', []), list) else o.get('correo', '—')),
                    ("Persona Contacto:", o.get('contacto', '—')),
                    ("Tlf. Contacto:", o.get('telefono_contacto', '—'))
                ]
                
                # Calcular altura necesaria: Cabecera(10) + Espacio(2) + Detalles(8 cada uno) + Margen(5)
                valid_detalles = [v for l, v in detalles_obra if v and str(v).strip() != '—']
                needed_h = 10 + 2 + (len(valid_detalles) * 8) + 5
                
                if pdf.get_y() + needed_h > 275: # Umbral para página A4
                    pdf.add_page()

                pdf.set_font("helvetica", "B", 12)
                pdf.set_fill_color(240, 240, 240)
                pdf.cell(0, 10, f"  {o.get('nombre_obra', 'Obra')}", ln=True, fill=True)
                pdf.ln(2)
                
                for label, value in detalles_obra:
                    if value and str(value).strip() != '—':
                        pdf.set_font("helvetica", "B", 10)
                        pdf.cell(40, 8, label)
                        pdf.set_font("helvetica", "", 10)
                        pdf.cell(0, 8, str(value), ln=True)
                pdf.ln(5)
            
        # Buffer
        output = io.BytesIO()
        pdf_bytes = pdf.output()
        output.write(pdf_bytes)
        output.seek(0)
        
        clean_name = secure_filename(casa.get('nombre', 'Casa'))
        filename = f"Reporte_{clean_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as exc:
        print(f"Error generando reporte de casa: {exc}")
        return "Error al generar el reporte", 500

@bp.route('/reporte_obra/<casa_id>/<obra_id>')
def reporte_obra(casa_id, obra_id):
    try:
        casa = MCasas.getCasaById(casa_id)
        if not casa:
            return "Casa no encontrada", 404
            
        obras = casa.get('obras', [])
        obra_data = next((o for o in obras if o.get('id') == obra_id), None)
        if not obra_data:
            return "Obra no encontrada", 404
            
        # Create PDF instance
        pdf = PDF()
        pdf.alias_nb_pages()
        pdf.add_page()
        
        # Report Date
        pdf.set_font("helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        now = datetime.now().strftime("%d/%m/%Y %H:%M")
        pdf.cell(0, 8, f"Ficha generada el: {now}", ln=True, align='R')
        pdf.ln(5)
        
        # Title
        pdf.set_font("helvetica", "B", 16)
        pdf.set_text_color(44, 62, 80)
        pdf.cell(0, 10, obra_data.get('nombre_obra', 'Detalle de la Obra'), ln=True, align='C')
        
        # Subtitle Casa Perteneciente
        casa_nombre = (casa.get('nombre') or '—') if casa else "—"
        pdf.cell(0, 8, f"Perteneciente a: {casa_nombre}", ln=True, align='C')
        pdf.ln(10)
        
        # Info sections
        pdf.set_font("helvetica", "B", 12)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(0, 10, "  Información General", ln=True, fill=True)
        pdf.ln(2)
        
        pdf.set_font("helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        
        telfs = obra_data.get('telefono', [])
        if isinstance(telfs, list):
            telf_str = ", ".join([str(t) for t in telfs])
        else:
            telf_str = str(telfs) if telfs else "—"

        details = [
            ("Ciudad:", obra_data.get('ciudad', '—')),
            ("Apartado Postal:", obra_data.get('apartado_postal', '—')),
            ("Teléfonos:", telf_str),
            ("Dirección:", obra_data.get('direccion', '—')),
            ("Sitio Web:", obra_data.get('web', '—')),
            ("Correo(s):", ", ".join(obra_data.get('correo', [])) if isinstance(obra_data.get('correo', []), list) else obra_data.get('correo', '—')),
        ]
        
        for label, value in details:
            if value and str(value).strip() != '—':
                pdf.set_font("helvetica", "B", 10)
                pdf.cell(40, 8, label)
                pdf.set_font("helvetica", "", 10)
                pdf.cell(0, 8, str(value), ln=True)
            
        pdf.ln(5)
        
        # Contact section
        pdf.set_font("helvetica", "B", 12)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(0, 10, "  Información de Contacto", ln=True, fill=True)
        pdf.ln(2)
        
        pdf.set_font("helvetica", "B", 10)
        pdf.cell(40, 8, "Persona:")
        pdf.set_font("helvetica", "", 10)
        pdf.cell(0, 8, str(obra_data.get('contacto', '—')), ln=True)
        
        pdf.set_font("helvetica", "B", 10)
        pdf.cell(40, 8, "Teléfono:")
        pdf.set_font("helvetica", "", 10)
        pdf.cell(0, 8, str(obra_data.get('telefono_contacto', '—')), ln=True)
        
        # Buffer
        output = io.BytesIO()
        pdf_bytes = pdf.output()
        output.write(pdf_bytes)
        output.seek(0)
        
        clean_name = secure_filename(obra_data.get('nombre_obra', 'Obra'))
        filename = f"Reporte_{clean_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as exc:
        print(f"Error generando reporte de obra: {exc}")
        return "Error al generar el reporte", 500

# ── FUNCIONES AUXILIARES WORD ───────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

# Colores corporativos (mismo que el PDF)
COLOR_RED    = RGBColor(0xDC, 0x1E, 0x46)   # #DC1E46
COLOR_NAVY   = RGBColor(0x2C, 0x3E, 0x50)   # #2C3E50
COLOR_LIGHT  = RGBColor(0xF0, 0xF0, 0xF0)   # #F0F0F0 (zebra par)
COLOR_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_DARK   = RGBColor(0x20, 0x20, 0x20)

def _set_cell_bg(cell, rgb_hex):
    """Pinta el fondo de una celda con color hexadecimal (ej. 'DC1E46')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)

def _cell_text(cell, text, bold=False, color=None, size_pt=10, align='left'):
    """Escribe texto en una celda con formato."""
    para = cell.paragraphs[0]
    para.clear()
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color

def _add_section_divider(doc, color_rgb='DC1E46'):
    """Añade una línea horizontal de color."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_rgb)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(0)

def _generar_word_base():
    doc = Document()

    # ── Márgenes ──
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Fuente por defecto ──
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── Encabezado de página ──
    section = doc.sections[0]
    header  = section.header

    # Tabla de encabezado: logo | institución | título derecho
    logo_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '..', '..', 'presentation', 'static', 'img', 'logo_sdb.png'
    )
    htbl = header.add_table(1, 3, width=Inches(6.8))
    htbl.columns[0].width = Inches(0.8)
    htbl.columns[1].width = Inches(3.5)
    htbl.columns[2].width = Inches(2.5)

    # Celda logo
    logo_cell = htbl.cell(0, 0)
    logo_p = logo_cell.paragraphs[0]
    if os.path.exists(logo_path):
        logo_p.add_run().add_picture(logo_path, width=Inches(0.6))

    # Celda institución
    inst_cell = htbl.cell(0, 1)
    inst_p = inst_cell.paragraphs[0]
    r1 = inst_p.add_run('CONDOR, La Vega R.D.\n')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = COLOR_NAVY
    r2 = inst_p.add_run('Diocesis de La Vega')
    r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Celda título
    tit_cell = htbl.cell(0, 2)
    tit_p = tit_cell.paragraphs[0]
    tit_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rt = tit_p.add_run('DIRECTORIO DE OBRAS')
    rt.bold = True; rt.font.size = Pt(13); rt.font.color.rgb = COLOR_RED

    # Línea roja bajo el encabezado
    _add_section_divider(header)

    # ── Pie de página ──
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('Agenda SDB — Sistema de Gestión')
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99); fr.italic = True

    return doc

def _tabla_casas(doc, obras, first_table=True):
    """Crea una tabla de obras con header rojo y zebra striping."""
    cols_labels = ['Nombre de la Obra', 'Ciudad', 'Teléfono', 'Correo(s)', 'Contacto']
    col_widths  = [Inches(1.9), Inches(1.1), Inches(1.3), Inches(1.6), Inches(1.5)]

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # ── Header row (rojo) ──
    hdr_row = table.rows[0]
    for i, (lbl, w) in enumerate(zip(cols_labels, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        _set_cell_bg(cell, 'DC1E46')
        _cell_text(cell, lbl, bold=True, color=COLOR_WHITE, size_pt=9, align='center')

    # ── Rows de datos ──
    if not obras:
        row = table.add_row()
        _set_cell_bg(row.cells[0], 'FFFFFF')
        row.cells[0].merge(row.cells[4])
        _cell_text(row.cells[0], '(Sin obras registradas)', color=COLOR_DARK, align='center')
    else:
        for idx, o in enumerate(obras):
            row = table.add_row()
            fill = 'F0F0F0' if idx % 2 == 1 else 'FFFFFF'

            telfs   = o.get('telefono', [])
            telf_str = ', '.join([str(t) for t in telfs]) if isinstance(telfs, list) else str(telfs or '—')

            correos = o.get('correo', [])
            email_str = ', '.join(correos) if isinstance(correos, list) else str(correos or '—')

            ciudad = o.get('ciudad') or '—'
            ap = o.get('apartado_postal')
            if ap: ciudad += f' (AP: {ap})'

            cont_val  = o.get('contacto') or '—'

            values = [
                str(o.get('nombre_obra') or '—'),
                ciudad,
                telf_str or '—',
                email_str or '—',
                cont_val,
            ]
            for i, val in enumerate(values):
                _set_cell_bg(row.cells[i], fill)
                _cell_text(row.cells[i], val, color=COLOR_DARK, size_pt=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

@bp.route('/reporte_casas_word')
def reporte_casas_word():
    try:
        casas_list_raw = MCasas.getAllCasas()
        grupos = {'masculino': [], 'femenino': []}
        for c in casas_list_raw:
            t = str(c.get('tipo', 'masculino')).lower()
            if t == 'femenino': grupos['femenino'].append(c)
            else: grupos['masculino'].append(c)

        labels = {
            'masculino': 'CASAS MASCULINAS (SDB)',
            'femenino':  'CASAS FEMENINAS (FMA / HH.SS.)'
        }

        doc = _generar_word_base()

        # Fecha de emisión
        now = datetime.now().strftime('%d/%m/%Y %H:%M')
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dr = dp.add_run(f'Fecha de emisión: {now}')
        dr.italic = True; dr.font.size = Pt(8); dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        first_group = True
        for tipo, casas in grupos.items():
            if not casas: continue

            if not first_group:
                doc.add_page_break()
            first_group = False

            # ── Título de sección ──
            sec_p = doc.add_paragraph()
            sec_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sec_p.paragraph_format.space_before = Pt(8)
            sec_p.paragraph_format.space_after  = Pt(4)
            sec_r = sec_p.add_run(labels.get(tipo, tipo.upper()))
            sec_r.bold = True; sec_r.font.size = Pt(16); sec_r.font.color.rgb = COLOR_RED
            _add_section_divider(doc)

            for c in casas:
                nombre_casa = c.get('nombre', '').strip()
                obras = c.get('obras', [])

                # ── Nombre de la Casa (encabezado azul marino) ──
                casa_p = doc.add_paragraph()
                casa_p.paragraph_format.space_before = Pt(10)
                casa_p.paragraph_format.space_after  = Pt(3)
                _set_cell_bg  # only for tables, use run color here
                casa_r = casa_p.add_run(f'  CASA: {nombre_casa.upper()}  ')
                casa_r.bold = True; casa_r.font.size = Pt(11); casa_r.font.color.rgb = COLOR_WHITE
                # Simulate dark background via XML shading on the paragraph
                pPr2 = casa_p._p.get_or_add_pPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), '2C3E50')
                pPr2.append(shd2)

                _tabla_casas(doc, obras)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        filename = f"Reporte_Casas_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)
    except Exception as exc:
        import traceback; traceback.print_exc()
        return str(exc), 500

@bp.route('/reporte_casa_word/<casa_id>')
def reporte_casa_word(casa_id):
    try:
        casa = MCasas.getCasaById(casa_id)
        if not casa: return 'No encontrada', 404

        doc = _generar_word_base()
        now = datetime.now().strftime('%d/%m/%Y %H:%M')

        # Fecha
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dr = dp.add_run(f'Ficha generada el: {now}')
        dr.italic = True; dr.font.size = Pt(8); dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Título (nombre de la casa)
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(6)
        tr = tp.add_run(casa.get('nombre', 'Detalle de la Casa'))
        tr.bold = True; tr.font.size = Pt(18); tr.font.color.rgb = COLOR_NAVY

        # Historia
        historia = casa.get('historia', '').strip()
        if historia:
            _add_section_divider(doc)
            hp = doc.add_paragraph()
            hr = hp.add_run('  Reseña Histórica  ')
            hr.bold = True; hr.font.size = Pt(12); hr.font.color.rgb = COLOR_WHITE
            hpPr = hp._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F0F0F0')
            hpPr.append(shd)
            hist_p = doc.add_paragraph(historia)
            hist_p.paragraph_format.space_before = Pt(4)
            hist_p.paragraph_format.left_indent  = Inches(0.2)
            hist_p.runs[0].font.size = Pt(10); hist_p.runs[0].font.color.rgb = COLOR_DARK

        # Obras
        obras = casa.get('obras', [])
        if obras:
            _add_section_divider(doc)
            op = doc.add_paragraph()
            or_ = op.add_run('  Obras Asociadas  ')
            or_.bold = True; or_.font.size = Pt(14); or_.font.color.rgb = COLOR_RED
            op.paragraph_format.space_before = Pt(8)

            for o in obras:
                obras_nombre_p = doc.add_paragraph()
                obras_nombre_p.paragraph_format.space_before = Pt(10)
                obras_nombre_p.paragraph_format.space_after  = Pt(3)
                on_r = obras_nombre_p.add_run(f'  {o.get("nombre_obra", "Obra")}  ')
                on_r.bold = True; on_r.font.size = Pt(11); on_r.font.color.rgb = COLOR_WHITE
                pPr3 = obras_nombre_p._p.get_or_add_pPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), '2C3E50')
                pPr3.append(shd3)

                correos = o.get('correo', [])
                email_str = ', '.join(correos) if isinstance(correos, list) else str(correos or '—')
                telfs = o.get('telefono', [])
                telf_str = ', '.join([str(t) for t in telfs]) if isinstance(telfs, list) else str(telfs or '—')

                details = [
                    ('Ciudad',           o.get('ciudad', '—')),
                    ('Teléfono',         telf_str or '—'),
                    ('Dirección',        o.get('direccion', '—')),
                    ('Sitio Web',        o.get('web', '—')),
                    ('Correo(s)',         email_str or '—'),
                    ('Persona Contacto', o.get('contacto', '—')),
                    ('Tlf. Contacto',    o.get('telefono_contacto', '—')),
                ]
                # Mini tabla de detalles (2 columnas: etiqueta | valor)
                dtbl = doc.add_table(rows=len(details), cols=2)
                dtbl.style = 'Table Grid'
                for i, (lbl, val) in enumerate(details):
                    fill_d = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
                    lbl_cell = dtbl.rows[i].cells[0]
                    val_cell = dtbl.rows[i].cells[1]
                    lbl_cell.width = Inches(1.5)
                    val_cell.width = Inches(5.7)
                    _set_cell_bg(lbl_cell, 'EAF0FB')
                    _set_cell_bg(val_cell, fill_d)
                    _cell_text(lbl_cell, lbl, bold=True, color=COLOR_NAVY, size_pt=9)
                    _cell_text(val_cell, val, color=COLOR_DARK, size_pt=9)
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        filename = f"Reporte_{secure_filename(casa.get('nombre','Casa'))}.docx"
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)
    except Exception as exc:
        import traceback; traceback.print_exc()
        return str(exc), 500

@bp.route('/reporte_obra_word/<casa_id>/<obra_id>')
def reporte_obra_word(casa_id, obra_id):
    try:
        casa = MCasas.getCasaById(casa_id)
        obra = next((o for o in casa.get('obras', []) if o.get('id') == obra_id), None)
        if not obra: return 'No encontrada', 404

        doc = _generar_word_base()
        now = datetime.now().strftime('%d/%m/%Y %H:%M')

        # Fecha
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dr = dp.add_run(f'Ficha generada el: {now}')
        dr.italic = True; dr.font.size = Pt(8); dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Título
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in [tp.add_run(obra.get('nombre_obra', 'Detalle de la Obra'))]:
            r.bold = True; r.font.size = Pt(18); r.font.color.rgb = COLOR_NAVY
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_r = sub_p.add_run(f'Perteneciente a: {casa.get("nombre", "—")}')
        sub_r.font.size = Pt(10); sub_r.font.color.rgb = RGBColor(0x60, 0x60, 0x60); sub_r.italic = True

        _add_section_divider(doc)

        correos = obra.get('correo', [])
        email_str = ', '.join(correos) if isinstance(correos, list) else str(correos or '—')
        telfs = obra.get('telefono', [])
        telf_str = ', '.join([str(t) for t in telfs]) if isinstance(telfs, list) else str(telfs or '—')

        # Sección: Información General
        sec1_p = doc.add_paragraph()
        sec1_r = sec1_p.add_run('  Información General  ')
        sec1_r.bold = True; sec1_r.font.size = Pt(12); sec1_r.font.color.rgb = COLOR_WHITE
        sp1 = sec1_p._p.get_or_add_pPr()
        shd1 = OxmlElement('w:shd')
        shd1.set(qn('w:val'), 'clear'); shd1.set(qn('w:color'), 'auto'); shd1.set(qn('w:fill'), 'F0F0F0')
        sp1.append(shd1)
        sec1_p.paragraph_format.space_before = Pt(8); sec1_p.paragraph_format.space_after = Pt(4)

        gen_details = [
            ('Ciudad',          obra.get('ciudad', '—')),
            ('Apartado Postal', obra.get('apartado_postal', '—')),
            ('Teléfonos',       telf_str or '—'),
            ('Dirección',       obra.get('direccion', '—')),
            ('Sitio Web',       obra.get('web', '—')),
            ('Correo(s)',        email_str or '—'),
        ]
        gtbl = doc.add_table(rows=len(gen_details), cols=2)
        gtbl.style = 'Table Grid'
        for i, (lbl, val) in enumerate(gen_details):
            fill_d = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
            lc = gtbl.rows[i].cells[0]; vc = gtbl.rows[i].cells[1]
            lc.width = Inches(1.5); vc.width = Inches(5.7)
            _set_cell_bg(lc, 'DDE3F0'); _set_cell_bg(vc, fill_d)
            _cell_text(lc, lbl, bold=True, color=COLOR_NAVY, size_pt=9)
            _cell_text(vc, val, color=COLOR_DARK, size_pt=9)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Sección: Contacto
        sec2_p = doc.add_paragraph()
        sec2_r = sec2_p.add_run('  Información de Contacto  ')
        sec2_r.bold = True; sec2_r.font.size = Pt(12); sec2_r.font.color.rgb = COLOR_WHITE
        sp2 = sec2_p._p.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), 'F0F0F0')
        sp2.append(shd2)
        sec2_p.paragraph_format.space_before = Pt(8); sec2_p.paragraph_format.space_after = Pt(4)

        cont_details = [
            ('Persona',    obra.get('contacto', '—')),
            ('Teléfono',   obra.get('telefono_contacto', '—')),
        ]
        ctbl = doc.add_table(rows=len(cont_details), cols=2)
        ctbl.style = 'Table Grid'
        for i, (lbl, val) in enumerate(cont_details):
            fill_d = 'F5F5F5' if i % 2 == 0 else 'FFFFFF'
            lc = ctbl.rows[i].cells[0]; vc = ctbl.rows[i].cells[1]
            lc.width = Inches(1.5); vc.width = Inches(5.7)
            _set_cell_bg(lc, 'DDE3F0'); _set_cell_bg(vc, fill_d)
            _cell_text(lc, lbl, bold=True, color=COLOR_NAVY, size_pt=9)
            _cell_text(vc, val, color=COLOR_DARK, size_pt=9)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        filename = f"Reporte_{secure_filename(obra.get('nombre_obra','Obra'))}.docx"
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)
    except Exception as exc:
        import traceback; traceback.print_exc()
        return str(exc), 500

